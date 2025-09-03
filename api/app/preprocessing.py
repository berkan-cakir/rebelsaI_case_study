import os
from datetime import datetime
from docx import Document
import re
from pathlib import Path

def get_file_metadata(file_path: str) -> dict:
    """Get basic metadata of a file: size in KB and last modified time."""
    try:
        stats = os.stat(file_path)

        return {
            "size_kb": stats.st_size / 1024,  # size in KB
            "modified_time": datetime.fromtimestamp(stats.st_mtime)
        }
    except Exception as e:
        print(f"Error getting metadata for {file_path}: {e}")
        return None, None
    
def process_folder(path: str) -> list:
    """Process all .docx files in the folder and return their metadata."""
    try:
        return [
            get_file_metadata(os.path.join(path, name))
            for name in os.listdir(path)
            if os.path.isfile(os.path.join(path, name)) and name.endswith('.docx')
        ]
    except Exception as e:
        print(f"Error processing folder {path}: {e}")
        return []
    
def clean_paragraph_text(text: str) -> str:
    """Basic cleaning: normalize whitespace, strip, remove weird chars."""
    text = text.replace("\xa0", " ")        # non-breaking spaces
    text = re.sub(r"\s+", " ", text)        # collapse multiple spaces/newlines
    return text.strip()
    
def extract_text_from_docx(file_path: str) -> str:
    """Extract and clean text from a .docx file, preserving some structure."""
    try:
        doc = Document(file_path)
        parts = []

        filename = Path(file_path).stem
        parts.append(f"# {filename}\n")  # Title as H1

        # Extract paragraphs in order
        for para in doc.paragraphs:
            text = clean_text(para.text)
            if text:
                style = para.style.name if para.style else ""
                if style.startswith("Heading"):
                    parts.append(f"\n## {text}\n")  # markdown-like heading H2
                elif "List" in style or text.startswith(("-", "•", "*", "1.", "a.")):
                    parts.append(f"- {text}")
                else:
                    parts.append(text)

        # Extract tables in order
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [clean_text(cell.text) for cell in row.cells]
                rows.append(cells)
            if rows:
                header = " | ".join(rows[0])
                sep = " | ".join(["---"] * len(rows[0]))
                body = [" | ".join(r) for r in rows[1:]]
                table_str = "\n".join([header, sep] + body)
                parts.append("\n" + table_str + "\n")

        # Join everything into one string
        return "\n".join(parts).strip()

    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return ""